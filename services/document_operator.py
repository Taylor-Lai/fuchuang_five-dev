"""文档操作执行服务"""
import os
from io import BytesIO
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.nlp_command_parser import OperationType, ParsedCommand


class DocumentOperator:
    """文档操作执行器"""
    
    def __init__(self, document_bytes: bytes):
        """
        初始化文档操作器
        
        Args:
            document_bytes: 文档的字节内容
        """
        self.doc = Document(BytesIO(document_bytes))
        self.original_bytes = document_bytes
    
    def execute_command(self, command: ParsedCommand) -> Dict:
        """
        执行解析后的指令
        
        Args:
            command: 解析后的指令对象
            
        Returns:
            Dict: 执行结果
        """
        try:
            # 获取操作参数
            params = command.params
            action = params.get("action", "")
            
            # 根据 LLM 的智能解析结果执行操作
            if action == "insert":
                result = self._smart_insert(params)
            elif action == "delete":
                result = self._smart_delete(params)
            elif action == "replace":
                result = self._smart_replace(params)
            elif action == "format":
                result = self._smart_format(params)
            elif action == "layout":
                result = self._smart_layout(params)
            elif action == "general":
                result = self._smart_general(params)
            else:
                # 尝试使用传统方法
                if command.operation_type == OperationType.EDIT:
                    result = self._execute_edit(params)
                elif command.operation_type == OperationType.FORMAT:
                    result = self._execute_format(params)
                elif command.operation_type == OperationType.LAYOUT:
                    result = self._execute_layout(params)
                elif command.operation_type == OperationType.EXTRACT:
                    result = self._execute_extract(params)
                elif command.operation_type == OperationType.STRUCTURE:
                    result = self._execute_structure(params)
                else:
                    return {"success": False, "message": "不支持的操作类型"}
            
            return {
                "success": True,
                "message": "操作执行成功",
                "result": result,
                "confidence": command.confidence
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"操作执行失败: {str(e)}",
                "confidence": command.confidence
            }
    
    def _smart_insert(self, params: Dict) -> Dict:
        """
        智能插入操作
        支持LLM生成的复杂插入指令
        """
        content = params.get("content", "")
        position = params.get("position", "end")
        
        # 检查是否有段落数量要求
        paragraph_count = params.get("paragraph_count", 0)
        
        # 如果内容为空或需要生成，使用LLM生成
        if not content or "随便写" in content or "自己写" in content or "生成" in content:
            document_context = "\n".join([p.text for p in self.doc.paragraphs[:5]])
            
            # 构建生成提示词
            if paragraph_count > 0:
                generation_prompt = f"请根据文档主题生成{paragraph_count}段相关内容，每段内容要独立成段，内容要连贯、有逻辑性。"
            else:
                # 检测是否包含段落数量要求
                import re
                match = re.search(r'(\d+)段', params.get('content', ''))
                if match:
                    para_count = int(match.group(1))
                    generation_prompt = f"请根据文档主题生成{para_count}段相关内容，每段内容要独立成段，内容要连贯、有逻辑性。"
                else:
                    generation_prompt = f"请根据文档主题生成相关内容，{params.get('details', {}).get('step1', '生成连贯的内容')}"
            
            # 生成内容
            generated_content = DocumentOperator.generate_content(generation_prompt, document_context)
            
            if generated_content:
                content = generated_content
            else:
                # 兜底内容
                if paragraph_count == 3:
                    content = "这是第一段内容。\n这是第二段内容。\n这是第三段内容。"
                elif paragraph_count == 5:
                    content = "这是第一段内容。\n这是第二段内容。\n这是第三段内容。\n这是第四段内容。\n这是第五段内容。"
                else:
                    content = "这是自动生成的内容。\n这是第二段内容。\n这是第三段内容。"
        
        # 处理多段内容
        # 首先按常见的段落分隔符分割
        paragraphs = []
        # 先按换行符分割
        lines = content.split('\n')
        for line in lines:
            # 去除前后空白
            line = line.strip()
            # 跳过空行
            if line:
                paragraphs.append(line)
        
        # 如果没有分割到段落，尝试按句号分割
        if len(paragraphs) == 1 and len(paragraphs[0]) > 100:
            sentences = paragraphs[0].split('。')
            paragraphs = []
            for sentence in sentences:
                sentence = sentence.strip()
                if sentence:
                    paragraphs.append(sentence + '。')
        
        # 确保至少有内容
        if not paragraphs:
            return {"message": "没有要插入的内容"}
        
        # 根据位置插入
        if position == "start":
            for para_text in reversed(paragraphs):
                self.doc.paragraphs[0].insert_paragraph_before(para_text)
        elif position == "end":
            for para_text in paragraphs:
                self.doc.add_paragraph(para_text)
        elif position == "after_paragraph" and "paragraph_index" in params:
            para_index = params["paragraph_index"]
            if para_index < len(self.doc.paragraphs):
                for para_text in reversed(paragraphs):
                    self.doc.paragraphs[para_index].insert_paragraph_before(para_text)
        else:
            # 默认在末尾插入
            for para_text in paragraphs:
                self.doc.add_paragraph(para_text)
        
        return {"message": f"已插入{len(paragraphs)}段内容"}
    
    def _smart_delete(self, params: Dict) -> Dict:
        """
        智能删除操作
        """
        target = params.get("target", "")
        
        if not target:
            return {"message": "未指定删除目标"}
        
        count = 0
        for para in self.doc.paragraphs:
            if target in para.text:
                # 只删除目标内容，而不是整个段落
                para.text = para.text.replace(target, "")
                count += 1
        
        return {"message": f"已删除{count}处内容"}
    
    def _smart_replace(self, params: Dict) -> Dict:
        """
        智能替换操作
        """
        old_text = params.get("target", "")
        new_text = params.get("content", "")
        
        if not old_text:
            return {"message": "未指定替换目标"}
        
        # 如果新文本需要生成
        if not new_text or "随便写" in new_text or "自己写" in new_text:
            document_context = "\n".join([p.text for p in self.doc.paragraphs[:5]])
            generation_prompt = f"请生成替换'{old_text}'的相关内容，保持文档风格一致"
            new_text = DocumentOperator.generate_content(generation_prompt, document_context)
        
        count = 0
        for para in self.doc.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
                count += 1
        
        return {"message": f"已替换{count}处内容"}
    
    def _smart_format(self, params: Dict) -> Dict:
        """
        智能格式操作
        """
        target = params.get("target", "")
        
        count = 0
        for para in self.doc.paragraphs:
            if target and target not in para.text:
                continue
            
            for run in para.runs:
                # 应用各种格式参数
                for key, value in params.items():
                    if key == "font_size":
                        run.font.size = Pt(value)
                    elif key == "font_color":
                        color_map = {
                            "red": RGBColor(255, 0, 0),
                            "blue": RGBColor(0, 0, 255),
                            "green": RGBColor(0, 128, 0),
                            "black": RGBColor(0, 0, 0),
                        }
                        if value.lower() in color_map:
                            run.font.color.rgb = color_map[value.lower()]
                    elif key == "bold":
                        run.font.bold = value
                    elif key == "italic":
                        run.font.italic = value
                count += 1
        
        return {"message": f"已修改{count}处格式"}
    
    def _smart_layout(self, params: Dict) -> Dict:
        """
        智能排版操作
        """
        target = params.get("target", "")
        
        count = 0
        for para in self.doc.paragraphs:
            if target and target not in para.text:
                continue
            
            # 应用各种排版参数
            for key, value in params.items():
                if key == "alignment":
                    alignment_map = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if value.lower() in alignment_map:
                        para.alignment = alignment_map[value.lower()]
                elif key == "indent":
                    para.paragraph_format.first_line_indent = Inches(value / 72.0)
            count += 1
        
        return {"message": f"已修改{count}处排版"}
    
    def _smart_general(self, params: Dict) -> Dict:
        """
        智能通用操作
        处理LLM生成的复杂操作指令
        """
        # 检查是否有详细步骤
        details = params.get("details", {})
        
        if details:
            # 执行多步骤操作
            results = []
            for step_key, step_value in details.items():
                if "插入" in step_value or "添加" in step_value:
                    result = self._smart_insert({"content": step_value, "position": "end"})
                    results.append(result)
            
            return {"message": f"已执行{len(results)}个步骤"}
        
        # 默认插入操作
        content = params.get("content", "")
        if content:
            return self._smart_insert({"content": content, "position": "end"})
        
        return {"message": "未找到可执行的操作"}
    
    def get_modified_document(self) -> bytes:
        """
        获取修改后的文档字节内容
        """
        output = BytesIO()
        self.doc.save(output)
        return output.getvalue()
    
    def get_document_preview(self) -> str:
        """
        获取文档预览文本
        """
        return "\n".join([p.text for p in self.doc.paragraphs[:10]])
    
    # 传统执行方法（保持向后兼容）
    def _execute_edit(self, params: Dict) -> Dict:
        """
        执行编辑操作
        """
        action = params.get("action", "")
        
        if action == "insert":
            return self._insert_text(params)
        elif action == "delete":
            return self._delete_text(params)
        elif action == "replace":
            return self._replace_text(params)
        else:
            return {"message": "不支持的编辑操作"}
    
    def _insert_text(self, params: Dict) -> Dict:
        """
        插入文本
        """
        content = params.get("content", "")
        position = params.get("position", "end")  # start, end, after_paragraph
        
        if position == "start":
            self.doc.paragraphs[0].insert_paragraph_before(content)
        elif position == "end":
            self.doc.add_paragraph(content)
        elif position == "after_paragraph" and "paragraph_index" in params:
            para_index = params["paragraph_index"]
            if para_index < len(self.doc.paragraphs):
                self.doc.paragraphs[para_index].insert_paragraph_before(content)
        
        return {"message": f"已插入文本: {content[:50]}..."}
    
    def _delete_text(self, params: Dict) -> Dict:
        """
        删除文本
        """
        target = params.get("target", "")
        paragraph_index = params.get("paragraph_index", -1)
        
        if paragraph_index >= 0 and paragraph_index < len(self.doc.paragraphs):
            self.doc.paragraphs[paragraph_index].clear()
            return {"message": f"已删除第{paragraph_index + 1}段"}
        elif target:
            # 根据内容查找并删除
            for i, para in enumerate(self.doc.paragraphs):
                if target in para.text:
                    para.clear()
                    return {"message": f"已删除包含'{target}'的段落"}
        
        return {"message": "未找到要删除的内容"}
    
    def _replace_text(self, params: Dict) -> Dict:
        """
        替换文本
        """
        old_text = params.get("target", "")
        new_text = params.get("content", "")
        
        if not old_text or not new_text:
            return {"message": "替换参数不完整"}
        
        count = 0
        for para in self.doc.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
                count += 1
        
        # 同时处理表格中的文本
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)
                        count += 1
        
        return {"message": f"已替换{count}处文本"}
    
    def _execute_format(self, params: Dict) -> Dict:
        """
        执行格式操作
        """
        target = params.get("target", "")
        font_size = params.get("font_size")
        font_color = params.get("font_color")
        bold = params.get("bold")
        italic = params.get("italic")
        
        count = 0
        # 处理段落格式
        for para in self.doc.paragraphs:
            if target and target not in para.text:
                continue
            
            for run in para.runs:
                if font_size:
                    run.font.size = Pt(font_size)
                if font_color:
                    color_map = {
                        "red": RGBColor(255, 0, 0),
                        "blue": RGBColor(0, 0, 255),
                        "green": RGBColor(0, 128, 0),
                        "black": RGBColor(0, 0, 0),
                    }
                    if font_color.lower() in color_map:
                        run.font.color.rgb = color_map[font_color.lower()]
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                count += 1
        
        return {"message": f"已修改{count}处格式"}
    
    def _execute_layout(self, params: Dict) -> Dict:
        """
        执行排版操作
        """
        target = params.get("target", "")
        alignment = params.get("alignment")
        indent = params.get("indent")
        
        count = 0
        for para in self.doc.paragraphs:
            if target and target not in para.text:
                continue
            
            if alignment:
                alignment_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if alignment.lower() in alignment_map:
                    para.alignment = alignment_map[alignment.lower()]
                    count += 1
            
            if indent:
                para.paragraph_format.first_line_indent = Inches(indent / 72.0)  # 转换为英寸
                count += 1
        
        return {"message": f"已修改{count}处排版"}
    
    def _execute_extract(self, params: Dict) -> Dict:
        """
        执行提取操作
        """
        extract_type = params.get("type", "text")
        criteria = params.get("criteria", "")
        
        results = []
        
        if extract_type == "text":
            for para in self.doc.paragraphs:
                if not criteria or criteria in para.text:
                    results.append(para.text)
        
        elif extract_type == "table":
            for i, table in enumerate(self.doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                results.append({"table_index": i, "data": table_data})
        
        return {
            "type": extract_type,
            "count": len(results),
            "data": results[:10]  # 限制返回数量
        }
    
    def _execute_structure(self, params: Dict) -> Dict:
        """
        执行结构操作
        """
        action = params.get("action", "")
        content = params.get("content", "")
        
        if action == "add_toc":
            # 添加目录（简单实现）
            self.doc.add_paragraph("目录", style="Heading 1")
            return {"message": "已添加目录"}
        elif action == "add_header":
            # 添加页眉
            section = self.doc.sections[0]
            header = section.header
            header.paragraphs[0].text = content
            return {"message": f"已添加页眉: {content}"}
        elif action == "add_section":
            # 添加章节
            self.doc.add_heading(content, level=1)
            return {"message": f"已添加章节: {content}"}
        else:
            return {"message": "不支持的结构操作"}