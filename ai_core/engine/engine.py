# any2table_engine/engine.py
import traceback
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from langchain_core.prompts import ChatPromptTemplate

from .schemas import *
from typing import Callable

from any2table.app import build_orchestrator
from any2table.config import AppConfig
from any2table.cli import discover_assets

# 加载环境变量
from dotenv import load_dotenv
load_dotenv()

# 选择 LLM 提供商
LLM_PROVIDER = os.getenv("LLM_PROVIDER", "zhipu")
ZHIPU_API_KEY = os.getenv("ZHIPU_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# 导入相应的 LLM 客户端
if LLM_PROVIDER == "zhipu" and ZHIPU_API_KEY:
    from langchain_community.chat_models import ChatZhipuAI
elif OPENAI_API_KEY:
    from langchain_openai import ChatOpenAI
else:
    raise ValueError("请在 .env 文件中配置 ZHIPU_API_KEY 或 OPENAI_API_KEY")


class FormatAction(BaseModel):
    target_paragraph_index: int = Field(..., description="要修改的段落索引(从0开始，如果是全文则填 -1)")
    font_size: Optional[int] = Field(None, description="字号大小(数字，如 14)")
    bold: Optional[bool] = Field(None, description="是否加粗")
    color_hex: Optional[str] = Field(None, description="十六进制颜色码，如 '#FF0000' 代表红色")
    alignment: Optional[str] = Field(None, description="对齐方式: 'left', 'center', 'right'")


class FormatPlan(BaseModel):
    actions: list[FormatAction] = Field(..., description="格式修改动作列表")


def handle_module_1_format(input_data: Mod1_FormatInput) -> Mod1_FormatOutput:
    try:
        doc_path = Path(input_data.file_path)
        doc = Document(doc_path)

        preview_text = "\n".join([f"[{i}] {p.text}" for i, p in enumerate(doc.paragraphs[:10]) if p.text.strip()])

        # 初始化 LLM 客户端
        if LLM_PROVIDER == "zhipu" and ZHIPU_API_KEY:
            llm = ChatZhipuAI(model="glm-4", api_key=ZHIPU_API_KEY, temperature=0)
        else:
            llm = ChatOpenAI(model="gpt-4o-mini", temperature=0, api_key=OPENAI_API_KEY)
        structured_llm = llm.with_structured_output(FormatPlan)

        prompt = ChatPromptTemplate.from_messages([
            ("system", "你是一个文档排版助手。以下是文档的前几段预览：\n{preview}\n\n请根据用户的要求，输出格式修改动作。"),
            ("human", "用户要求：{command}")
        ])

        plan: FormatPlan = structured_llm.invoke(
            prompt.format(preview=preview_text, command=input_data.natural_language_cmd))

        # 检查 plan 是否为 None
        if plan is None:
            return Mod1_FormatOutput(status="failed", message="AI 未能生成有效的格式修改计划，请检查 API 密钥或重试")

        # 检查 actions 是否为 None 或空
        if plan.actions is None or len(plan.actions) == 0:
            return Mod1_FormatOutput(status="failed", message="AI 未能生成任何格式修改动作，请尝试更具体的指令")

        for action in plan.actions:
            # 检查段落索引是否有效
            if action.target_paragraph_index == -1:
                paras_to_modify = doc.paragraphs
            else:
                # 确保索引在有效范围内
                if 0 <= action.target_paragraph_index < len(doc.paragraphs):
                    paras_to_modify = [doc.paragraphs[action.target_paragraph_index]]
                else:
                    # 索引无效，跳过该动作
                    continue

            for p in paras_to_modify:
                if action.alignment == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif action.alignment == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                for run in p.runs:
                    if action.bold is not None: run.bold = action.bold
                    if action.font_size is not None: run.font.size = Pt(action.font_size)
                    if action.color_hex:
                        hex_color = action.color_hex.lstrip('#')
                        run.font.color.rgb = RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16),
                                                      int(hex_color[4:], 16))

        # 4. 保存新文件
        output_path = doc_path.parent / f"{doc_path.stem}_formatted{doc_path.suffix}"
        doc.save(output_path)

        return Mod1_FormatOutput(status="success", processed_file_path=str(output_path), message="排版修改成功")

    except Exception as e:
        return Mod1_FormatOutput(status="failed", message=traceback.format_exc())


def handle_module_2_extract(input_data: Mod2_ExtractInput) -> Mod2_ExtractOutput:
    try:
        # 1. 读取文件文本
        file_path = input_data.file_path
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif file_path.endswith('.xlsx'):
            import pandas as pd
            df = pd.read_excel(file_path)
            full_text = df.to_markdown(index=False)
        else:
            # 尝试多种编码读取文本文件
            full_text = None
            for encoding in ['utf-8', 'gbk', 'utf-16']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        full_text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if full_text is None:
                return Mod2_ExtractOutput(status="failed", message=f"无法识别文件编码: {file_path}")

        from pydantic import create_model
        fields_spec = {entity: (str, Field(default="未找到", description=f"提取 '{entity}' 的内容")) for entity in
                      input_data.target_entities}
        DynamicExtractionModel = create_model('DynamicExtractionModel', **fields_spec)

        # 初始化 LLM 客户端
        if LLM_PROVIDER == "zhipu" and ZHIPU_API_KEY:
            llm = ChatZhipuAI(model="glm-4", api_key=ZHIPU_API_KEY, temperature=0)
        else:
            llm = ChatOpenAI(model="gpt-4o-mini", temperature=0, api_key=OPENAI_API_KEY)
        structured_llm = llm.with_structured_output(DynamicExtractionModel)

        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "你是一个精准的信息提取 AI。请从以下文本中提取指定字段。如果没有找到，请填'未找到'。\n\n文本内容：\n{text}"),
            ("human", "请提取以下字段：{entities}")
        ])

        result = structured_llm.invoke(
            prompt.format(text=full_text[:8000], entities=", ".join(input_data.target_entities)))

        # 检查 result 是否为 None
        if result is None:
            return Mod2_ExtractOutput(status="failed", message="AI 未能生成有效的提取结果，请检查 API 密钥或重试")

        return Mod2_ExtractOutput(status="success", extracted_data=result.dict())

    except Exception as e:
        return Mod2_ExtractOutput(status="failed", message=traceback.format_exc())


def handle_module_3_fusion(
        input_data: Mod3_FusionInput,
        progress_callback: Callable[[str, str, str], None] = None
) -> Mod3_FusionOutput:
    """
    执行多智能体数据融合任务。
    :param input_data: 包含 task_id, workspace_dir, user_request 的输入参数
    :param progress_callback: 进度回调函数，格式为 callback(task_id, status, message)
    """
    work_dir = Path(input_data.workspace_dir)

    # 确保总是创建用户要求文件，即使用户没有提供
    user_request_content = input_data.user_request if input_data.user_request else "请根据源文档中的信息，自动填充模板表格。"
    (work_dir / "用户要求.txt").write_text(user_request_content, encoding="utf-8")

    try:
        if progress_callback:
            progress_callback(input_data.task_id, "processing", "正在初始化多智能体融合环境...")

        # 配置 LLM
        llm_provider = LLM_PROVIDER
        llm_model = "glm-4" if llm_provider == "zhipu" else "gpt-4o-mini"
        llm_api_key_env = "ZHIPU_API_KEY" if llm_provider == "zhipu" else "OPENAI_API_KEY"
        
        config = AppConfig(
            enable_agent_runtime=True,
            agent_runtime_backend="langgraph",
            enable_llm_skill_execution=True,
            llm_provider=llm_provider,
            llm_model=llm_model,
            llm_api_key_env=llm_api_key_env,
        )

        assets = discover_assets(work_dir)

        if progress_callback:
            progress_callback(input_data.task_id, "processing",
                              f"已识别 {len(assets)} 个文档资产，启动 LangGraph 引擎...")

        orchestrator = build_orchestrator(config=config)

        if progress_callback:
            progress_callback(input_data.task_id, "processing", "7大智能体开始协同流转，进行跨文件特征提取与对齐...")

        result = orchestrator.run(assets)

        if progress_callback:
            progress_callback(input_data.task_id, "success", "多源数据融合与质检完成，物理资产已生成！")

        return Mod3_FusionOutput(
            status="success",
            task_id=input_data.task_id,
            output_excel_path=str(result.fill_result.output_path),
            warnings=result.fill_result.warnings
        )

    except Exception as e:
        if progress_callback:
            progress_callback(input_data.task_id, "failed", f"流转异常打断: {str(e)}")

        return Mod3_FusionOutput(
            status="failed",
            task_id=input_data.task_id,
            error_msg=traceback.format_exc()
        )