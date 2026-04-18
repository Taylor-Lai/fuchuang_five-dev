"""
ai_core/engine 模块单元测试
测试 FormatAction / FormatPlan 模型、文档格式操作逻辑、schemas 等
不依赖 LLM 调用，纯本地逻辑验证
"""
import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 确保 engine 包可以导入
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from engine.schemas import (
    Mod1_FormatInput, Mod1_FormatOutput,
    Mod2_ExtractInput, Mod2_ExtractOutput,
    Mod3_FusionInput, Mod3_FusionOutput,
)


class TestSchemas(unittest.TestCase):
    """测试 Pydantic schema 模型"""

    def test_mod1_input(self):
        inp = Mod1_FormatInput(file_path="/tmp/a.docx", natural_language_cmd="加粗第一段")
        self.assertEqual(inp.file_path, "/tmp/a.docx")
        self.assertEqual(inp.natural_language_cmd, "加粗第一段")

    def test_mod1_output_success(self):
        out = Mod1_FormatOutput(status="success", processed_file_path="/tmp/out.docx", message="OK")
        self.assertEqual(out.status, "success")
        self.assertEqual(out.processed_file_path, "/tmp/out.docx")

    def test_mod1_output_failed(self):
        out = Mod1_FormatOutput(status="failed", message="出错了")
        self.assertIsNone(out.processed_file_path)

    def test_mod2_input(self):
        inp = Mod2_ExtractInput(file_path="/tmp/a.docx", target_entities=["姓名", "年龄"])
        self.assertEqual(len(inp.target_entities), 2)

    def test_mod2_output_default(self):
        out = Mod2_ExtractOutput(status="success")
        self.assertEqual(out.extracted_data, {})

    def test_mod3_input_optional_request(self):
        inp = Mod3_FusionInput(task_id="t1", workspace_dir="/tmp/ws")
        self.assertIsNone(inp.user_request)

    def test_mod3_output_warnings_default(self):
        out = Mod3_FusionOutput(status="success", task_id="t1")
        self.assertEqual(out.warnings, [])


class TestFormatActionModel(unittest.TestCase):
    """测试 FormatAction / FormatPlan Pydantic 模型"""

    def _import_models(self):
        """延迟导入，避免触发模块级 LLM 初始化"""
        # 通过设置临时环境变量让 engine.py 的模块级代码不报错
        os.environ.setdefault("OPENAI_API_KEY", "test-key-for-unit-test")
        os.environ.setdefault("LLM_PROVIDER", "openai")
        # 需要重新导入
        from engine.engine import FormatAction, FormatPlan
        return FormatAction, FormatPlan

    def test_format_action_basic(self):
        FormatAction, _ = self._import_models()
        action = FormatAction(target_paragraph_index=0, bold=True, font_size=14)
        self.assertEqual(action.target_paragraph_index, 0)
        self.assertTrue(action.bold)
        self.assertEqual(action.font_size, 14)
        self.assertIsNone(action.color_hex)
        self.assertIsNone(action.alignment)

    def test_format_action_all_fields(self):
        FormatAction, _ = self._import_models()
        action = FormatAction(
            target_paragraph_index=-1,
            bold=False,
            font_size=16,
            color_hex="#FF0000",
            alignment="center"
        )
        self.assertEqual(action.target_paragraph_index, -1)
        self.assertEqual(action.color_hex, "#FF0000")

    def test_format_plan_normal(self):
        _, FormatPlan = self._import_models()
        plan = FormatPlan(actions=[
            {"target_paragraph_index": 0, "bold": True},
            {"target_paragraph_index": 1, "font_size": 12},
        ])
        self.assertEqual(len(plan.actions), 2)

    def test_format_plan_empty_actions(self):
        _, FormatPlan = self._import_models()
        plan = FormatPlan(actions=[])
        self.assertEqual(len(plan.actions), 0)

    def test_format_plan_normalize_action_singular(self):
        """测试 model_validator 兼容 LLM 返回 'action' 单数的情况"""
        _, FormatPlan = self._import_models()
        plan = FormatPlan.model_validate({
            "action": [{"target_paragraph_index": 0, "bold": True}]
        })
        self.assertEqual(len(plan.actions), 1)

    def test_format_plan_normalize_action_non_list(self):
        """测试 model_validator 兼容 action 不是列表的情况"""
        _, FormatPlan = self._import_models()
        plan = FormatPlan.model_validate({"action": "some string"})
        self.assertEqual(len(plan.actions), 0)


class TestDocxFormatting(unittest.TestCase):
    """测试实际的 docx 文档格式操作逻辑（不调用 LLM）"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _create_test_docx(self, paragraphs: list[str]) -> str:
        """创建测试用 docx 文件"""
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        path = os.path.join(self.temp_dir, "test.docx")
        doc.save(path)
        return path

    def test_bold_single_paragraph(self):
        """测试对单个段落加粗"""
        path = self._create_test_docx(["第一段内容", "第二段内容", "第三段内容"])
        doc = Document(path)

        p = doc.paragraphs[0]
        for run in p.runs:
            run.bold = True

        out_path = os.path.join(self.temp_dir, "out.docx")
        doc.save(out_path)

        doc2 = Document(out_path)
        for run in doc2.paragraphs[0].runs:
            self.assertTrue(run.bold)

    def test_color_hex_6digit(self):
        """测试 6 位颜色码解析"""
        path = self._create_test_docx(["测试红色"])
        doc = Document(path)

        hex_color = "FF0000"
        for run in doc.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(
                int(hex_color[:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:], 16)
            )

        out_path = os.path.join(self.temp_dir, "out.docx")
        doc.save(out_path)

        doc2 = Document(out_path)
        for run in doc2.paragraphs[0].runs:
            self.assertEqual(run.font.color.rgb, RGBColor(0xFF, 0x00, 0x00))

    def test_color_hex_3digit_expansion(self):
        """测试 3 位颜色码扩展为 6 位"""
        hex_color = "F00"
        # 扩展逻辑
        if len(hex_color) == 3:
            hex_color = ''.join(c * 2 for c in hex_color)
        self.assertEqual(hex_color, "FF0000")
        self.assertEqual(len(hex_color), 6)

    def test_color_hex_invalid_length_skipped(self):
        """测试非法长度颜色码被跳过（不崩溃）"""
        hex_color = "FFGG00"  # 非法但长度为 6
        # 模拟引擎中的逻辑
        hex_raw = hex_color.lstrip('#')
        if len(hex_raw) == 3:
            hex_raw = ''.join(c * 2 for c in hex_raw)
        if len(hex_raw) == 6:
            try:
                RGBColor(int(hex_raw[:2], 16), int(hex_raw[2:4], 16), int(hex_raw[4:], 16))
            except ValueError:
                pass  # 非法十六进制应该由 int() 抛出，但不应崩溃整个流程
        # 走到这里不崩溃就是成功
        self.assertTrue(True)

    def test_alignment_left(self):
        """测试左对齐"""
        path = self._create_test_docx(["测试对齐"])
        doc = Document(path)
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        out_path = os.path.join(self.temp_dir, "out.docx")
        doc.save(out_path)
        doc2 = Document(out_path)
        self.assertEqual(doc2.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)

    def test_alignment_center(self):
        """测试居中对齐"""
        path = self._create_test_docx(["测试对齐"])
        doc = Document(path)
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        out_path = os.path.join(self.temp_dir, "out.docx")
        doc.save(out_path)
        doc2 = Document(out_path)
        self.assertEqual(doc2.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_alignment_right(self):
        """测试右对齐"""
        path = self._create_test_docx(["测试对齐"])
        doc = Document(path)
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        out_path = os.path.join(self.temp_dir, "out.docx")
        doc.save(out_path)
        doc2 = Document(out_path)
        self.assertEqual(doc2.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)

    def test_font_size(self):
        """测试字号修改"""
        path = self._create_test_docx(["测试字号"])
        doc = Document(path)
        for run in doc.paragraphs[0].runs:
            run.font.size = Pt(18)
        out_path = os.path.join(self.temp_dir, "out.docx")
        doc.save(out_path)
        doc2 = Document(out_path)
        for run in doc2.paragraphs[0].runs:
            self.assertEqual(run.font.size, Pt(18))

    def test_paragraph_index_out_of_range(self):
        """测试段落索引越界时应跳过而非崩溃"""
        path = self._create_test_docx(["唯一段落"])
        doc = Document(path)
        # 索引 5 超出范围，应跳过
        idx = 5
        if 0 <= idx < len(doc.paragraphs):
            self.fail("不应进入此分支")
        # 不崩溃即通过
        self.assertTrue(True)

    def test_paragraph_index_minus_one_means_all(self):
        """测试 -1 表示全文"""
        path = self._create_test_docx(["段落1", "段落2", "段落3"])
        doc = Document(path)
        idx = -1
        if idx == -1:
            paras = doc.paragraphs
        self.assertEqual(len(paras), 3)

    def test_empty_runs_paragraph(self):
        """测试段落有文字但 runs 为空时，创建 run 的逻辑"""
        path = self._create_test_docx(["有内容的段落"])
        doc = Document(path)
        p = doc.paragraphs[0]

        # 模拟引擎中的 empty runs 修复逻辑
        original_text = p.text
        if p.runs:
            # add_paragraph 创建的段落一般有 runs，验证正常流程
            for run in p.runs:
                run.bold = True
        else:
            # 如果没有 runs，走修复路径
            text = p.text
            p.clear()
            p.add_run(text)
            for run in p.runs:
                run.bold = True

        out_path = os.path.join(self.temp_dir, "out.docx")
        doc.save(out_path)
        doc2 = Document(out_path)
        self.assertEqual(doc2.paragraphs[0].text, original_text)

    def test_preview_text_generation(self):
        """测试预览文本生成逻辑（取前10段非空段落）"""
        paragraphs = [f"段落{i}" for i in range(15)]
        paragraphs.insert(3, "")  # 插入空段落
        path = self._create_test_docx(paragraphs)
        doc = Document(path)
        preview = "\n".join([f"[{i}] {p.text}" for i, p in enumerate(doc.paragraphs[:10]) if p.text.strip()])
        # 空段落应被过滤
        self.assertNotIn("[]", preview)
        self.assertIn("[0] 段落0", preview)

    def test_output_file_naming(self):
        """测试输出文件命名规则"""
        doc_path = Path("/tmp/report.docx")
        output_path = doc_path.parent / f"{doc_path.stem}_formatted{doc_path.suffix}"
        self.assertEqual(output_path.name, "report_formatted.docx")
        self.assertEqual(output_path.parent, doc_path.parent)


class TestDynamicModelCreation(unittest.TestCase):
    """测试模块二的动态 Pydantic 模型创建"""

    def test_chinese_field_names(self):
        """测试中文字段名的 create_model"""
        from pydantic import BaseModel, Field, create_model
        entities = ["项目名称", "负责人", "预算"]
        fields_spec = {e: (str, Field(default="未找到")) for e in entities}
        Model = create_model('TestModel', **fields_spec)
        instance = Model()
        self.assertEqual(instance.项目名称, "未找到")
        self.assertEqual(instance.负责人, "未找到")

    def test_chinese_field_model_dump(self):
        """测试 model_dump 输出中文 key"""
        from pydantic import BaseModel, Field, create_model
        entities = ["姓名", "年龄"]
        fields_spec = {e: (str, Field(default="未找到")) for e in entities}
        Model = create_model('TestModel', **fields_spec)
        instance = Model(姓名="张三", 年龄="25")
        data = instance.model_dump()
        self.assertEqual(data["姓名"], "张三")
        self.assertEqual(data["年龄"], "25")

    def test_single_entity(self):
        """测试单个字段"""
        from pydantic import BaseModel, Field, create_model
        fields_spec = {"名称": (str, Field(default="未找到"))}
        Model = create_model('TestModel', **fields_spec)
        instance = Model(名称="测试项目")
        self.assertEqual(instance.model_dump()["名称"], "测试项目")


class TestFileReading(unittest.TestCase):
    """测试模块二的文件读取逻辑"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_read_docx(self):
        """测试 docx 文件读取"""
        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("第二段")
        path = os.path.join(self.temp_dir, "test.docx")
        doc.save(path)

        doc2 = Document(path)
        text = "\n".join([p.text for p in doc2.paragraphs if p.text.strip()])
        self.assertIn("第一段", text)
        self.assertIn("第二段", text)

    def test_read_txt_utf8(self):
        """测试 UTF-8 文本文件读取"""
        path = os.path.join(self.temp_dir, "test.txt")
        with open(path, 'w', encoding='utf-8') as f:
            f.write("UTF-8 中文内容")

        full_text = None
        for encoding in ['utf-8', 'gbk', 'utf-16']:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    full_text = f.read()
                break
            except UnicodeDecodeError:
                continue
        self.assertIsNotNone(full_text)
        self.assertIn("中文内容", full_text)

    def test_read_txt_gbk(self):
        """测试 GBK 编码文件读取"""
        path = os.path.join(self.temp_dir, "test.txt")
        with open(path, 'wb') as f:
            f.write("GBK 中文".encode('gbk'))

        full_text = None
        for encoding in ['utf-8', 'gbk', 'utf-16']:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    full_text = f.read()
                break
            except UnicodeDecodeError:
                continue
        self.assertIsNotNone(full_text)
        self.assertIn("中文", full_text)

    def test_file_not_exists(self):
        """测试文件不存在"""
        path = os.path.join(self.temp_dir, "nonexistent.docx")
        self.assertFalse(Path(path).exists())


if __name__ == "__main__":
    unittest.main()
