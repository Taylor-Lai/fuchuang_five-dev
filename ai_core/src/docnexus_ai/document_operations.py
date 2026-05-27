"""Natural-language Word document operations."""

from __future__ import annotations

import traceback
import json
from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field, model_validator

from docnexus_ai.llm import get_chat_llm

def _schema_classes():
    try:
        from ai_core.engine.schemas import Mod1_FormatOutput
    except ImportError:  # pragma: no cover - compatibility for tests importing engine directly
        from engine.schemas import Mod1_FormatOutput
    return Mod1_FormatOutput


class FormatAction(BaseModel):
    operation: str = Field("format", description="操作类型: format, insert, delete, replace, extract, structure")
    target_paragraph_index: int = Field(-1, description="要修改的段落索引(从0开始，如果是全文则填 -1)")
    target_text: str | None = Field(None, description="要查找或操作的目标文本")
    content: str | None = Field(None, description="插入内容、替换内容或结构内容")
    font_size: int | None = Field(None, description="字号大小(数字，如 14)")
    bold: bool | None = Field(None, description="是否加粗")
    color_hex: str | None = Field(None, description="十六进制颜色码，如 '#FF0000' 代表红色")
    alignment: str | None = Field(None, description="对齐方式: 'left', 'center', 'right'")


class FormatPlan(BaseModel):
    actions: list[FormatAction] = Field(default_factory=list, description="格式修改动作列表，若无需修改则为空列表")

    @model_validator(mode="before")
    @classmethod
    def _normalize_actions(cls, data):
        if isinstance(data, dict) and "action" in data and "actions" not in data:
            val = data["action"]
            data["actions"] = val if isinstance(val, list) else []
        return data


PARAGRAPH_INDEX_WORDS = {
    "一": 0,
    "二": 1,
    "三": 2,
    "四": 3,
    "五": 4,
    "六": 5,
    "七": 6,
    "八": 7,
    "九": 8,
    "十": 9,
}
COLOR_MAP = {
    "红": "#FF0000",
    "红色": "#FF0000",
    "蓝": "#0000FF",
    "蓝色": "#0000FF",
    "绿": "#008000",
    "绿色": "#008000",
    "黑": "#000000",
    "黑色": "#000000",
}


def _split_command(command: str) -> list[str]:
    parts = re.split(r"[；;。\n]+|然后|并且|同时|最后|接着|随后|再", command)
    return [part.strip(" ，,") for part in parts if part.strip(" ，,")]


def _infer_target_paragraph_index(text: str) -> int | None:
    if any(token in text for token in ("全文", "全部", "所有")):
        return -1
    digit_match = re.search(r"第\s*(\d+)\s*段", text)
    if digit_match:
        return max(int(digit_match.group(1)) - 1, 0)
    cn_match = re.search(r"第\s*([一二三四五六七八九十])\s*段", text)
    if cn_match:
        return PARAGRAPH_INDEX_WORDS.get(cn_match.group(1), -1)
    return None


def _infer_alignment(text: str) -> str | None:
    if "居中" in text or "居中对齐" in text:
        return "center"
    if "右对齐" in text or "靠右" in text:
        return "right"
    if "左对齐" in text or "靠左" in text:
        return "left"
    return None


def _infer_color(text: str) -> str | None:
    hex_match = re.search(r"#?[0-9a-fA-F]{6}", text)
    if hex_match:
        color = hex_match.group(0)
        return color if color.startswith("#") else f"#{color}"
    for token, color in COLOR_MAP.items():
        if token in text:
            return color
    return None


def _infer_font_size(text: str) -> int | None:
    match = re.search(r"(\d+)\s*(?:号|pt|磅)", text, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def _extract_quoted_pair(text: str) -> tuple[str | None, str | None]:
    quoted = re.findall(r"[“\"']([^”\"']+)[”\"']", text)
    if len(quoted) >= 2:
        return quoted[0], quoted[1]
    return None, None


def build_rule_based_plan(command: str) -> FormatPlan:
    actions: list[FormatAction] = []
    last_target_index = -1
    for part in _split_command(command):
        explicit_target_index = _infer_target_paragraph_index(part)
        target_index = explicit_target_index if explicit_target_index is not None else last_target_index
        if explicit_target_index is not None:
            last_target_index = explicit_target_index
        if "替换" in part or "改成" in part or "改为" in part:
            old_text, new_text = _extract_quoted_pair(part)
            if old_text is None:
                match = re.search(r"(?:把|将)?(.+?)(?:替换为|替换成|改成|改为)(.+)", part)
                if match:
                    old_text = match.group(1).strip(" ，,。")
                    new_text = match.group(2).strip(" ，,。")
            if old_text and new_text:
                actions.append(FormatAction(operation="replace", target_paragraph_index=target_index, target_text=old_text, content=new_text))
                continue
        if "删除" in part or "去掉" in part:
            target_text = None
            quoted = re.findall(r"[“\"']([^”\"']+)[”\"']", part)
            if quoted:
                target_text = quoted[0]
            actions.append(FormatAction(operation="delete", target_paragraph_index=target_index, target_text=target_text))
            continue
        if "插入" in part or "添加" in part or "新增" in part:
            quoted = re.findall(r"[“\"']([^”\"']+)[”\"']", part)
            content = quoted[0] if quoted else None
            if content:
                if "开头" in part or "前面" in part:
                    target_index = 0
                elif "末尾" in part or "最后" in part:
                    target_index = -2
                actions.append(FormatAction(operation="insert", target_paragraph_index=target_index, content=content))
                continue
        if "目录" in part:
            actions.append(FormatAction(operation="structure", target_paragraph_index=-1, target_text="目录"))
            continue
        if "页眉" in part:
            quoted = re.findall(r"[“\"']([^”\"']+)[”\"']", part)
            actions.append(FormatAction(operation="structure", target_paragraph_index=-1, target_text="页眉", content=quoted[0] if quoted else None))
            continue
        if "提取" in part or "抽取" in part:
            actions.append(FormatAction(operation="extract", target_paragraph_index=target_index))
            continue

        bold = True if "加粗" in part else None
        font_size = _infer_font_size(part)
        color_hex = _infer_color(part)
        alignment = _infer_alignment(part)
        if any(value is not None for value in (bold, font_size, color_hex, alignment)):
            new_action = FormatAction(
                operation="format",
                target_paragraph_index=target_index,
                bold=bold,
                font_size=font_size,
                color_hex=color_hex,
                alignment=alignment,
            )
            if actions and actions[-1].operation == "format" and actions[-1].target_paragraph_index == target_index:
                previous = actions[-1]
                previous.bold = new_action.bold if new_action.bold is not None else previous.bold
                previous.font_size = new_action.font_size if new_action.font_size is not None else previous.font_size
                previous.color_hex = new_action.color_hex if new_action.color_hex is not None else previous.color_hex
                previous.alignment = new_action.alignment if new_action.alignment is not None else previous.alignment
            else:
                actions.append(new_action)
    return FormatPlan(actions=actions)


def _target_paragraphs(doc: Document, action: FormatAction):
    if action.target_paragraph_index == -1:
        return doc.paragraphs
    elif 0 <= action.target_paragraph_index < len(doc.paragraphs):
        return [doc.paragraphs[action.target_paragraph_index]]
    elif action.target_text:
        return [p for p in doc.paragraphs if action.target_text in p.text]
    return []


def _apply_format_action(doc: Document, action: FormatAction) -> int:
    paragraphs = _target_paragraphs(doc, action)
    changed = 0

    for paragraph in paragraphs:
        if action.alignment == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif action.alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif action.alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if not paragraph.runs and paragraph.text:
            text = paragraph.text
            paragraph.clear()
            paragraph.add_run(text)

        for run in paragraph.runs:
            if action.bold is not None:
                run.bold = action.bold
            if action.font_size is not None:
                run.font.size = Pt(action.font_size)
            if action.color_hex:
                hex_color = action.color_hex.lstrip("#")
                if len(hex_color) == 3:
                    hex_color = "".join(c * 2 for c in hex_color)
                if len(hex_color) == 6:
                    run.font.color.rgb = RGBColor(
                        int(hex_color[:2], 16),
                        int(hex_color[2:4], 16),
                        int(hex_color[4:], 16),
                    )
            changed += 1
    return changed


def _apply_insert_action(doc: Document, action: FormatAction) -> int:
    content = (action.content or "").strip()
    if not content:
        return 0
    paragraphs = [line.strip() for line in content.splitlines() if line.strip()] or [content]
    if action.target_paragraph_index == -2:
        for text in paragraphs:
            doc.add_paragraph(text)
    elif action.target_paragraph_index == 0 and doc.paragraphs:
        for text in reversed(paragraphs):
            doc.paragraphs[0].insert_paragraph_before(text)
    elif 0 <= action.target_paragraph_index < len(doc.paragraphs):
        anchor = doc.paragraphs[action.target_paragraph_index]
        for text in reversed(paragraphs):
            anchor.insert_paragraph_before(text)
    else:
        for text in paragraphs:
            doc.add_paragraph(text)
    return len(paragraphs)


def _apply_delete_action(doc: Document, action: FormatAction) -> int:
    changed = 0
    target = action.target_text
    for paragraph in _target_paragraphs(doc, action):
        if target:
            if target in paragraph.text:
                paragraph.text = paragraph.text.replace(target, "")
                changed += 1
        else:
            paragraph.clear()
            changed += 1
    return changed


def _apply_replace_action(doc: Document, action: FormatAction) -> int:
    target = action.target_text or ""
    content = action.content or ""
    if not target:
        return 0
    changed = 0
    for paragraph in doc.paragraphs:
        if target in paragraph.text:
            paragraph.text = paragraph.text.replace(target, content)
            changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if target in cell.text:
                    cell.text = cell.text.replace(target, content)
                    changed += 1
    return changed


def _apply_structure_action(doc: Document, action: FormatAction) -> int:
    content = (action.content or "").strip()
    target = (action.target_text or "").strip()
    if target in {"目录", "toc"}:
        doc.add_paragraph("目录", style="Heading 1")
        return 1
    if target in {"标题", "heading"} and content:
        doc.add_heading(content, level=1)
        return 1
    if target in {"页眉", "header"} and content:
        section = doc.sections[0]
        section.header.paragraphs[0].text = content
        return 1
    return 0


def _apply_extract_action(doc: Document, action: FormatAction) -> tuple[int, list[str]]:
    paragraphs = _target_paragraphs(doc, action)
    extracted = [paragraph.text for paragraph in paragraphs if paragraph.text.strip()]
    return len(extracted), extracted


def _execute_action(doc: Document, action: FormatAction) -> tuple[str, int, dict[str, object]]:
    operation = (action.operation or "format").lower()
    if operation == "insert":
        changed = _apply_insert_action(doc, action)
        return operation, changed, {}
    if operation == "delete":
        changed = _apply_delete_action(doc, action)
        return operation, changed, {}
    if operation == "replace":
        changed = _apply_replace_action(doc, action)
        return operation, changed, {}
    if operation == "structure":
        changed = _apply_structure_action(doc, action)
        return operation, changed, {}
    if operation == "extract":
        changed, extracted = _apply_extract_action(doc, action)
        return operation, changed, {"extracted_text": extracted[:20]}
    changed = _apply_format_action(doc, action)
    return "format", changed, {}


def _write_operation_audit(
    doc_path: Path,
    output_path: Path,
    command: str,
    actions: list[FormatAction],
    action_reports: list[dict[str, object]],
    warnings: list[str] | None = None,
) -> Path:
    audit_path = output_path.with_suffix(".operation_audit.json")
    payload = {
        "schema_version": "1.0",
        "source_file": str(doc_path),
        "output_file": str(output_path),
        "command": command,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "action_count": len(actions),
        "actions": action_reports,
        "missed_actions": [report for report in action_reports if report["affected_count"] == 0],
        "warnings": list(warnings or []),
    }
    audit_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit_path


def handle_document_operation(input_data):
    Mod1_FormatOutput = _schema_classes()
    try:
        doc_path = Path(input_data.file_path)
        if not doc_path.exists():
            return Mod1_FormatOutput(status="failed", message=f"文件不存在: {doc_path}")

        doc = Document(doc_path)
        preview_text = "\n".join(
            f"[{i}] {p.text}" for i, p in enumerate(doc.paragraphs[:10]) if p.text.strip()
        )

        rule_plan = build_rule_based_plan(input_data.natural_language_cmd)
        prompt = ChatPromptTemplate.from_messages([
            ("system", (
                "你是一个文档智能操作助手。以下是文档的前几段预览：\n{preview}\n\n"
                "请根据用户要求输出可执行动作列表，支持 format/insert/delete/replace/extract/structure。\n"
                "输出格式示例：\n"
                '{{"actions": [{{"operation": "format", "target_paragraph_index": 0, "bold": true, "font_size": 16, "alignment": "center", "color_hex": null}}]}}\n'
                "字段说明：\n"
                "- operation: format/insert/delete/replace/extract/structure\n"
                "- target_paragraph_index: 段落索引（从0开始），-1 表示全文\n"
                "- target_text: 删除、替换、结构操作的目标文本或目标类型\n"
                "- content: 插入内容、替换内容、标题/页眉内容\n"
                "- bold: true/false/null\n"
                "- font_size: 整数字号或 null\n"
                "- alignment: 'left'/'center'/'right' 或 null\n"
                "- color_hex: 十六进制颜色如 '#FF0000' 或 null\n"
                "不需要修改的字段填 null，必须至少返回一个 action。"
            )),
            ("human", "用户要求：{command}"),
        ])

        plan = None
        warnings: list[str] = []
        try:
            structured_llm = get_chat_llm().with_structured_output(FormatPlan)
            plan = (prompt | structured_llm).invoke({
                "preview": preview_text,
                "command": input_data.natural_language_cmd,
            })
        except Exception as exc:
            warnings.append(f"LLM plan generation failed; used rule-based fallback: {exc}")
            plan = None
        if plan is None:
            plan = rule_plan
        elif rule_plan.actions:
            # Keep deterministic actions as a safety net for multi-intent commands.
            existing = {(action.operation, action.target_paragraph_index, action.target_text, action.content) for action in plan.actions}
            for action in rule_plan.actions:
                key = (action.operation, action.target_paragraph_index, action.target_text, action.content)
                if key not in existing:
                    plan.actions.append(action)
        if not plan.actions:
            return Mod1_FormatOutput(status="failed", message="AI 未能生成任何格式修改动作，请尝试更具体的指令")

        action_summaries = []
        action_reports = []
        for action in plan.actions:
            operation, changed, extra = _execute_action(doc, action)
            action_summaries.append(f"{operation}:{changed}")
            action_reports.append({
                "operation": operation,
                "affected_count": changed,
                "target_paragraph_index": action.target_paragraph_index,
                "target_text": action.target_text,
                "content_preview": (action.content or "")[:120],
                **extra,
            })

        output_path = doc_path.parent / f"{doc_path.stem}_formatted{doc_path.suffix}"
        doc.save(output_path)
        audit_path = _write_operation_audit(
            doc_path,
            output_path,
            input_data.natural_language_cmd,
            plan.actions,
            action_reports,
            warnings,
        )
        warning_text = f"；警告：{'；'.join(warnings)}" if warnings else ""
        return Mod1_FormatOutput(
            status="success",
            processed_file_path=str(output_path),
            message=f"文档智能操作完成，动作统计：{', '.join(action_summaries)}；审计报告：{audit_path}{warning_text}",
        )

    except Exception:
        return Mod1_FormatOutput(status="failed", message=traceback.format_exc())
