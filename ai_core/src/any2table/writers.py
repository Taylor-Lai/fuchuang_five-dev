"""Writers."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from any2table.core.models import CanonicalDocument, CellWriteTrace, FillResult, StructuredRecord, TemplateSpec

logger = logging.getLogger(__name__)


def _ensure_output_path(template_doc: CanonicalDocument) -> Path:
    src_path = Path(template_doc.file.path)
    output_dir = src_path.parent / "outputs"
    output_dir.mkdir(exist_ok=True)
    return output_dir / f"{src_path.stem}-filled{src_path.suffix}"


def _build_fallback_output_path(output_path: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    return output_path.with_name(f"{output_path.stem}-{timestamp}{output_path.suffix}")


def _save_with_fallback(save_func, output_path: Path) -> tuple[Path, list[str]]:
    try:
        save_func(output_path)
        return output_path, []
    except PermissionError:
        fallback_path = _build_fallback_output_path(output_path)
        save_func(fallback_path)
        return fallback_path, [f"Primary output path was locked, wrote fallback file: {fallback_path}"]


class XlsxWriter:
    supported_doc_types = ("xlsx",)

    def write(
        self,
        template_doc: CanonicalDocument,
        template_spec: TemplateSpec,
        records: list[StructuredRecord],
    ) -> FillResult:
        output_path = _ensure_output_path(template_doc)
        workbook = load_workbook(template_doc.file.path)
        written_cells: list[CellWriteTrace] = []
        records_by_table: dict[str, list[StructuredRecord]] = {}
        for record in records:
            records_by_table.setdefault(record.target_table_id, []).append(record)

        for table_index, target_table in enumerate(template_spec.target_tables):
            if table_index >= len(workbook.worksheets):
                logger.warning(
                    "Target table %s (index %d) has no corresponding worksheet; skipping",
                    target_table.target_table_id, table_index,
                )
                continue
            worksheet = workbook.worksheets[table_index]
            data_start_row = 2
            target_records = records_by_table.get(target_table.target_table_id, [])
            for row_offset, record in enumerate(target_records):
                row_number = data_start_row + row_offset
                for col_index, field in enumerate(target_table.schema, start=1):
                    value = record.values.get(field.field_name)
                    worksheet.cell(row=row_number, column=col_index, value=value)
                    written_cells.append(
                        CellWriteTrace(
                            target_table_id=target_table.target_table_id,
                            row_index=row_number - 1,
                            col_index=col_index - 1,
                            field_name=field.field_name,
                            value=value,
                            record_id=record.record_id,
                            evidence_ids=record.field_sources.get(field.field_name, []),
                        )
                    )

        final_output_path, warnings = _save_with_fallback(workbook.save, output_path)
        return FillResult(
            output_doc_id=template_doc.doc_id,
            output_path=str(final_output_path),
            written_cells=written_cells,
            warnings=warnings,
        )


class DocxTableWriter:
    supported_doc_types = ("docx",)

    def write(
        self,
        template_doc: CanonicalDocument,
        template_spec: TemplateSpec,
        records: list[StructuredRecord],
    ) -> FillResult:
        output_path = _ensure_output_path(template_doc)
        document = Document(template_doc.file.path)
        records_by_table: dict[str, list[StructuredRecord]] = {}
        for record in records:
            records_by_table.setdefault(record.target_table_id, []).append(record)

        written_cells: list[CellWriteTrace] = []
        inserted_rows: list[dict[str, object]] = []

        for table_index, target_table in enumerate(template_spec.target_tables):
            if table_index >= len(document.tables):
                logger.warning(
                    "Target table %s (index %d) has no corresponding docx table; skipping",
                    target_table.target_table_id, table_index,
                )
                continue
            docx_table = document.tables[table_index]
            target_records = records_by_table.get(target_table.target_table_id, [])
            if not target_records:
                continue

            existing_data_rows = max(len(docx_table.rows) - 1, 0)
            while existing_data_rows < len(target_records):
                docx_table.add_row()
                existing_data_rows += 1
                inserted_rows.append({"target_table_id": target_table.target_table_id, "row_index": existing_data_rows})

            for record_index, record in enumerate(target_records, start=1):
                if record_index >= len(docx_table.rows):
                    logger.warning(
                        "Not enough rows in docx table for target %s; stopping at record %d of %d",
                        target_table.target_table_id, record_index, len(target_records),
                    )
                    break
                row = docx_table.rows[record_index]
                for col_index, field in enumerate(target_table.schema):
                    if col_index >= len(row.cells):
                        continue
                    value = record.values.get(field.field_name)
                    row.cells[col_index].text = "" if value is None else str(value)
                    written_cells.append(
                        CellWriteTrace(
                            target_table_id=target_table.target_table_id,
                            row_index=record_index,
                            col_index=col_index,
                            field_name=field.field_name,
                            value=value,
                            record_id=record.record_id,
                            evidence_ids=record.field_sources.get(field.field_name, []),
                        )
                    )

        final_output_path, warnings = _save_with_fallback(document.save, output_path)
        return FillResult(
            output_doc_id=template_doc.doc_id,
            output_path=str(final_output_path),
            written_cells=written_cells,
            inserted_rows=inserted_rows,
            warnings=warnings,
        )
